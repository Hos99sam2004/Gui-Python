#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Document sections creation functions
"""
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils import (
    shade_cell, add_decorative_line, add_section_heading, add_bullet_point
)


def add_header(doc):
    """Add document header with title and decorations"""
    add_decorative_line(doc)
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("📋 تقرير المشروع - تطبيق إدارة الهوية")
    title_run.font.size = Pt(26)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(25, 118, 210)
    
    # Subtitle with project info
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Identity Management Application")
    subtitle_run.font.size = Pt(11)
    subtitle_run.font.italic = True
    subtitle_run.font.color.rgb = RGBColor(100, 100, 100)
    
    # Date and info
    from datetime import datetime
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"📅 التاريخ: {datetime.now().strftime('%d/%m/%Y')}  |  🔧 الإصدار: 1.0")
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = RGBColor(80, 80, 80)
    
    add_decorative_line(doc)
    doc.add_paragraph()


def add_objectives_section(doc):
    """Add objectives section"""
    add_decorative_line(doc, color=(150, 150, 150), size=11)
    add_section_heading(doc, "🎯 1. أهداف المشروع")
    
    objectives = [
        "🔐 تطوير تطبيق واجهة رسومية (GUI) باستخدام PySide6 لإدارة الهوية والعناوين IP",
        "🌐 دمج شبكة Tor للحصول على عناوين IP مختلفة وتغيير الهوية بشكل آمن",
        "📝 تطبيق نظام تسجيل شامل يحفظ سجل التغييرات بصيغة JSON Lines",
        "🛡️ إضافة ميزات الأمان مثل التحقق من سلامة المواقع عبر VirusTotal API",
        "💬 تفعيل نظام الإخطارات عبر Telegram للتنبيهات الفورية",
        "✨ توفير واجهة رسومية سهلة الاستخدام وآمنة للمستخدمين"
    ]
    
    for obj in objectives:
        add_bullet_point(doc, obj)
    
    doc.add_paragraph()


def add_achievements_section(doc):
    """Add achievements section"""
    add_decorative_line(doc, color=(150, 150, 150), size=11)
    add_section_heading(doc, "✅ 2. الإنجازات المحققة")
    
    achievements = [
        ("🔐 نظام إدارة الهوية الأساسي", [
            "✓ الاتصال الآمن بشبكة Tor",
            "✓ الحصول على عنوان IP الحقيقي",
            "✓ الحصول على عنوان IP من خلال Tor",
            "✓ تغيير الهوية (Tor NEWNYM)",
            "✓ الحصول على بيانات الموقع الجغرافي"
        ]),
        ("🖥️ واجهة المستخدم الرسومية", [
            "✓ تصميم واجهة احترافية وسهلة الاستخدام",
            "✓ عرض معلومات IP والموقع الجغرافي في الوقت الفعلي",
            "✓ زر لتغيير الهوية مع مؤشرات التقدم",
            "✓ عرض سجل التغييرات السابقة",
            "✓ لوحة إعدادات متقدمة"
        ]),
        ("📊 نظام التسجيل والإشعارات", [
            "✓ تسجيل جميع التغييرات في ملف log",
            "✓ حفظ البيانات بصيغة JSON Lines",
            "✓ عرض السجل التاريخي في الواجهة",
            "✓ دعم إشعارات Telegram (معد)"
        ]),
        ("🛡️ ميزات الأمان", [
            "✓ التحقق من سلامة المواقع عبر VirusTotal API",
            "✓ كشف المواقع المريبة والخطيرة",
            "✓ حظر المواقع المشبوهة تلقائياً",
            "✓ تقارير أمان مفصلة"
        ]),
        ("⚙️ البنية التقنية", [
            "✓ استخدام dataclasses لإدارة البيانات",
            "✓ نظام threading للعمليات غير المتزامنة",
            "✓ معالجة شاملة للأخطاء والاستثناءات",
            "✓ دعم متعدد الأنظمة (Windows, Linux, macOS)"
        ])
    ]
    
    for section_title, items in achievements:
        p = doc.add_paragraph()
        p_run = p.add_run(section_title)
        p_run.bold = True
        p_run.font.size = Pt(12)
        p_run.font.color.rgb = RGBColor(25, 118, 210)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(6)
        
        for item in items:
            item_p = doc.add_paragraph(item, style='List Bullet')
            item_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            item_p.paragraph_format.space_after = Pt(4)
            for run in item_p.runs:
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(60, 60, 60)
    
    doc.add_paragraph()


def add_progress_section(doc):
    """Add progress table section"""
    add_decorative_line(doc, color=(150, 150, 150), size=11)
    add_section_heading(doc, "📊 3. نسبة إنجاز المشروع")
    
    # Create progress table
    table = doc.add_table(rows=7, cols=3)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Header row with styling
    header_cells = table.rows[0].cells
    header_labels = ["النسبة", "الحالة", "المكون"]
    for i, label in enumerate(header_labels):
        cell = header_cells[i]
        cell.text = label
        shade_cell(cell, "1976D2")
        
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(255, 255, 255)
    
    # Data rows
    data = [
        ("95%", "✅ مكتمل تماماً", "نظام إدارة الهوية والعمليات الأساسية"),
        ("90%", "✅ مكتمل تماماً", "واجهة المستخدم الرسومية"),
        ("85%", "✅ مكتمل تماماً", "نظام التسجيل والسجلات"),
        ("70%", "🔄 قيد التطوير", "التكامل مع VirusTotal وفحص المواقع"),
        ("60%", "🔄 قيد التطوير", "نظام الإخطارات عبر Telegram"),
        ("80%", "⭐ مكتمل تقريباً", "المشروع الكلي")
    ]
    
    colors = ["E8F5E9", "FFF3E0", "FFF3E0", "E3F2FD", "E3F2FD", "F3E5F5"]
    
    for idx, (percentage, status, component) in enumerate(data, start=1):
        cells = table.rows[idx].cells
        cells[0].text = percentage
        cells[1].text = status
        cells[2].text = component
        
        for i, cell in enumerate(cells):
            shade_cell(cell, colors[idx-1])
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)
                run.font.bold = (i == 0)
    
    doc.add_paragraph()


def add_summary_section(doc):
    """Add summary and results section"""
    add_decorative_line(doc, color=(150, 150, 150), size=11)
    add_section_heading(doc, "📈 4. الخلاصة والنتائج")
    
    summary_points = [
        "🎯 نسبة الإنجاز الإجمالية للمشروع: 80% من المتطلبات الأساسية",
        "✨ تم تطوير جميع الميزات الأساسية بنجاح",
        "🎨 الواجهة الرسومية جاهزة للاستخدام وتوفر تجربة مستخدم احترافية",
        "🔒 نظام الأمان يعمل بكفاءة مع دعم كامل للمصادقة عبر Tor",
        "📋 سجل التغييرات يوفر تتبعاً كاملاً لجميع العمليات",
        "🚀 المشروع جاهز للنشر والاستخدام مع إمكانية إضافة ميزات إضافية في المستقبل"
    ]
    
    for point in summary_points:
        add_bullet_point(doc, point)
    
    doc.add_paragraph()


def add_next_steps_section(doc):
    """Add next steps and improvements section"""
    add_decorative_line(doc, color=(150, 150, 150), size=11)
    add_section_heading(doc, "🔮 5. الخطوات التالية والتحسينات المقترحة")
    
    next_steps = [
        "🔍 تكامل كامل مع VirusTotal API لفحص جميع المواقع",
        "🔔 تفعيل نظام الإخطارات عبر Telegram بالكامل",
        "🌐 إضافة دعم VPN بجانب Tor",
        "💻 تطوير نسخة ويب للتطبيق",
        "💾 إضافة ميزات نسخ احتياطي واستعادة البيانات",
        "⚡ تحسين الأداء وتقليل استهلاك الموارد",
        "🧪 إضافة اختبارات وحدة شاملة",
        "📚 توثيق شامل للمتطورين (Developer Documentation)"
    ]
    
    for step in next_steps:
        p = doc.add_paragraph(step, style='List Bullet')
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(60, 60, 60)
    
    doc.add_paragraph()
    doc.add_paragraph()


def add_footer(doc):
    """Add document footer"""
    add_decorative_line(doc)
    
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("✨ تم إعداد هذا التقرير بنجاح ✨")
    footer_run.font.size = Pt(11)
    footer_run.font.bold = True
    footer_run.font.color.rgb = RGBColor(25, 118, 210)
    
    from datetime import datetime
    footer2 = doc.add_paragraph()
    footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer2_run = footer2.add_run(f"تاريخ الإعداد: {datetime.now().strftime('%d/%m/%Y - %H:%M:%S')}")
    footer2_run.font.size = Pt(9)
    footer2_run.font.italic = True
    footer2_run.font.color.rgb = RGBColor(120, 120, 120)
