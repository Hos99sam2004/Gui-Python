#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

def shade_cell(cell, fill_color):
    """Add background color to cell"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._element.get_or_add_tcPr().append(shading_elm)

# Create document
doc = Document()

# Set document margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

# Header with decorative line
header_line = doc.add_paragraph()
header_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
header_run = header_line.add_run("━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
header_run.font.size = Pt(12)
header_run.font.color.rgb = RGBColor(25, 118, 210)

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run("PROJECT REPORT")
title_run.font.size = Pt(28)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(25, 118, 210)

# Subtitle
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.add_run("Identity Management Application with Tor Integration")
subtitle_run.font.size = Pt(13)
subtitle_run.font.italic = True
subtitle_run.font.color.rgb = RGBColor(80, 80, 80)

# Project Info
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info_run = info.add_run(f"Date: {datetime.now().strftime('%d %B %Y')} | Version: 1.0 | Status: 80% Complete")
info_run.font.size = Pt(10)
info_run.font.color.rgb = RGBColor(100, 100, 100)

# Header footer line
footer_line = doc.add_paragraph()
footer_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer_line.add_run("━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
footer_run.font.size = Pt(12)
footer_run.font.color.rgb = RGBColor(25, 118, 210)

doc.add_paragraph()  # Space

# ===== EXECUTIVE SUMMARY =====
exec_summary = doc.add_heading("Executive Summary", level=1)
exec_summary.alignment = WD_ALIGN_PARAGRAPH.LEFT
for run in exec_summary.runs:
    run.font.color.rgb = RGBColor(25, 118, 210)

exec_text = doc.add_paragraph(
    "This report documents the development and progress of a comprehensive Identity Management Application "
    "utilizing Tor network integration. The application provides a secure, user-friendly graphical interface "
    "for managing digital identities and IP addresses. As of the reporting date, the project has achieved "
    "80% completion of core functionality, with all primary objectives successfully implemented."
)
exec_text.alignment = WD_ALIGN_PARAGRAPH.LEFT
for run in exec_text.runs:
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(50, 50, 50)

doc.add_paragraph()  # Space

# ===== Section 1: PROJECT OBJECTIVES =====
section_div1 = doc.add_paragraph()
section_div1.alignment = WD_ALIGN_PARAGRAPH.CENTER
div1_run = section_div1.add_run("━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
div1_run.font.size = Pt(10)
div1_run.font.color.rgb = RGBColor(150, 150, 150)

heading1 = doc.add_heading("1. Project Objectives", level=1)
heading1.alignment = WD_ALIGN_PARAGRAPH.LEFT
for run in heading1.runs:
    run.font.color.rgb = RGBColor(25, 118, 210)

objectives = [
    "Develop a comprehensive GUI application using PySide6 for secure identity and IP address management",
    "Integrate Tor network protocol to obtain varied IP addresses and facilitate anonymous identity changes",
    "Implement comprehensive logging system with JSON Lines format for complete change history tracking",
    "Deploy security features including URL reputation verification via VirusTotal API",
    "Enable real-time notifications through Telegram for immediate security alerts",
    "Provide an intuitive, secure, and professional user interface for end-users"
]

for i, obj in enumerate(objectives, 1):
    p = doc.add_paragraph(f"{i}. {obj}")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.3)
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(50, 50, 50)

doc.add_paragraph()  # Space

# ===== Section 2: ACHIEVEMENTS =====
section_div2 = doc.add_paragraph()
section_div2.alignment = WD_ALIGN_PARAGRAPH.CENTER
div2_run = section_div2.add_run("━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
div2_run.font.size = Pt(10)
div2_run.font.color.rgb = RGBColor(150, 150, 150)

heading2 = doc.add_heading("2. Achievements and Deliverables", level=1)
heading2.alignment = WD_ALIGN_PARAGRAPH.LEFT
for run in heading2.runs:
    run.font.color.rgb = RGBColor(25, 118, 210)

achievements = [
    ("Core Identity Management System", [
        "Secure Tor network integration and connectivity",
        "Real IP address retrieval functionality",
        "Tor exit node IP address acquisition",
        "Identity rotation capability (Tor NEWNYM protocol)",
        "Geolocation data acquisition for IP addresses"
    ]),
    ("User Interface", [
        "Professional and intuitive graphical interface design",
        "Real-time IP and geolocation information display",
        "Identity change functionality with progress indicators",
        "Historical change log viewer with detailed records",
        "Advanced configuration panel for system settings"
    ]),
    ("Logging and Audit Trail", [
        "Comprehensive change tracking and logging",
        "JSON Lines format data storage for easy parsing",
        "Historical log display within application interface",
        "Telegram notification integration (configured)"
    ]),
    ("Security Features", [
        "URL reputation verification via VirusTotal API",
        "Malicious and suspicious URL detection",
        "Automatic blocking of flagged websites",
        "Detailed security assessment reports"
    ]),
    ("Technical Architecture", [
        "Dataclasses implementation for data management",
        "Multi-threaded architecture for asynchronous operations",
        "Comprehensive error handling and exception management",
        "Cross-platform compatibility (Windows, Linux, macOS)"
    ])
]

for section_title, items in achievements:
    p = doc.add_paragraph()
    p_run = p.add_run(f"• {section_title}")
    p_run.bold = True
    p_run.font.size = Pt(12)
    p_run.font.color.rgb = RGBColor(25, 118, 210)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.2)
    
    for item in items:
        item_p = doc.add_paragraph(item)
        item_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        item_p.paragraph_format.space_after = Pt(3)
        item_p.paragraph_format.left_indent = Inches(0.5)
        for run in item_p.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(60, 60, 60)

doc.add_paragraph()  # Space

# ===== Section 3: PROJECT COMPLETION STATUS =====
section_div3 = doc.add_paragraph()
section_div3.alignment = WD_ALIGN_PARAGRAPH.CENTER
div3_run = section_div3.add_run("━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
div3_run.font.size = Pt(10)
div3_run.font.color.rgb = RGBColor(150, 150, 150)

heading3 = doc.add_heading("3. Project Completion Status", level=1)
heading3.alignment = WD_ALIGN_PARAGRAPH.LEFT
for run in heading3.runs:
    run.font.color.rgb = RGBColor(25, 118, 210)

# Create status table
table = doc.add_table(rows=7, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Header row
header_cells = table.rows[0].cells
headers = ["Component", "Status", "Completion"]
for i, header_text in enumerate(headers):
    cell = header_cells[i]
    cell.text = header_text
    shade_cell(cell, "1976D2")
    
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in para.runs:
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(255, 255, 255)

# Data rows
data = [
    ("Core Identity Management System", "Completed", "95%"),
    ("Graphical User Interface", "Completed", "90%"),
    ("Logging and Audit System", "Completed", "85%"),
    ("VirusTotal Integration", "In Progress", "70%"),
    ("Telegram Notifications", "In Progress", "60%"),
    ("Overall Project Status", "Final Phase", "80%")
]

colors = ["E8F5E9", "C8E6C9", "FFF3E0", "FFCC80", "E3F2FD", "F3E5F5"]

for idx, (component, status, completion) in enumerate(data, start=1):
    cells = table.rows[idx].cells
    
    cells[0].text = component
    cells[1].text = status
    cells[2].text = completion
    
    for i, cell in enumerate(cells):
        shade_cell(cell, colors[idx-1])
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(10)
            run.font.bold = (i == 0)

doc.add_paragraph()  # Space

# ===== Section 4: CONCLUSIONS =====
section_div4 = doc.add_paragraph()
section_div4.alignment = WD_ALIGN_PARAGRAPH.CENTER
div4_run = section_div4.add_run("━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
div4_run.font.size = Pt(10)
div4_run.font.color.rgb = RGBColor(150, 150, 150)

heading4 = doc.add_heading("4. Conclusions and Key Findings", level=1)
heading4.alignment = WD_ALIGN_PARAGRAPH.LEFT
for run in heading4.runs:
    run.font.color.rgb = RGBColor(25, 118, 210)

conclusions = [
    "The project has successfully achieved 80% completion of all core requirements",
    "All primary functionality has been implemented and tested successfully",
    "The user interface is production-ready and provides professional user experience",
    "Security systems are operational with full Tor authentication support",
    "Historical audit trail provides complete operational tracking capabilities",
    "The application is ready for deployment with potential for future enhancements"
]

for i, conclusion in enumerate(conclusions, 1):
    p = doc.add_paragraph(f"{i}. {conclusion}")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.3)
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(50, 50, 50)

doc.add_paragraph()  # Space

# ===== Section 5: RECOMMENDATIONS =====
section_div5 = doc.add_paragraph()
section_div5.alignment = WD_ALIGN_PARAGRAPH.CENTER
div5_run = section_div5.add_run("━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
div5_run.font.size = Pt(10)
div5_run.font.color.rgb = RGBColor(150, 150, 150)

heading5 = doc.add_heading("5. Recommendations and Future Enhancements", level=1)
heading5.alignment = WD_ALIGN_PARAGRAPH.LEFT
for run in heading5.runs:
    run.font.color.rgb = RGBColor(25, 118, 210)

recommendations = [
    "Complete VirusTotal API integration for comprehensive URL scanning",
    "Activate full Telegram notification system for real-time alerts",
    "Implement VPN support alongside Tor for additional anonymity options",
    "Develop web-based version of the application for broader accessibility",
    "Add data backup and recovery features for enhanced reliability",
    "Optimize performance and reduce resource consumption",
    "Establish comprehensive unit testing framework",
    "Create detailed technical documentation for developers"
]

for i, rec in enumerate(recommendations, 1):
    p = doc.add_paragraph(f"{i}. {rec}")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.3)
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(60, 60, 60)

doc.add_paragraph()
doc.add_paragraph()

# Footer decorative line
footer_line2 = doc.add_paragraph()
footer_line2.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_line2_run = footer_line2.add_run("━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
footer_line2_run.font.size = Pt(12)
footer_line2_run.font.color.rgb = RGBColor(25, 118, 210)

# Report metadata
metadata = doc.add_paragraph()
metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
metadata_run = metadata.add_run("Report Generated Successfully")
metadata_run.font.size = Pt(10)
metadata_run.font.bold = True
metadata_run.font.color.rgb = RGBColor(25, 118, 210)

timestamp = doc.add_paragraph()
timestamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
timestamp_run = timestamp.add_run(f"Report Date: {datetime.now().strftime('%B %d, %Y at %H:%M:%S')}")
timestamp_run.font.size = Pt(9)
timestamp_run.font.italic = True
timestamp_run.font.color.rgb = RGBColor(120, 120, 120)

# Save
doc.save('Project_Report_English.docx')
print("✓ Report successfully created: Project_Report_English.docx")
