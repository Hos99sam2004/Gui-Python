#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Utility functions for document creation
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def shade_cell(cell, fill_color):
    """Add background color to cell"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._element.get_or_add_tcPr().append(shading_elm)


def setup_document_margins(doc, top=0.8, bottom=0.8, left=1, right=1):
    """Set document margins"""
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)


def add_decorative_line(doc, color=(25, 118, 210), size=14):
    """Add a decorative line separator"""
    line_para = doc.add_paragraph()
    line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line_run = line_para.add_run("━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
    line_run.font.size = Pt(size)
    line_run.font.color.rgb = RGBColor(*color)


def add_section_heading(doc, text, level=1, color=(25, 118, 210)):
    """Add a section heading with custom styling"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    heading_format = heading.paragraph_format
    heading_format.space_before = Pt(6)
    heading_format.space_after = Pt(6)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(*color)
    return heading


def add_bullet_point(doc, text, font_size=11, color=(50, 50, 50), rtl=True):
    """Add a bullet point with custom styling"""
    p = doc.add_paragraph(text, style='List Bullet')
    if rtl:
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(8)
    for run in p.runs:
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor(*color)
    return p
